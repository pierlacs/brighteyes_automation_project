#!/usr/bin/env python3
"""
Tender parser — Demo 1 (3C Card)

Goal
-----
Read a tender document (Markdown) plus the BrightEyes company context (Markdown),
extract key procurement structure (lots, eligibility, award criteria, deadlines, budget),
run a go/no-go check, match lots to BrightEyes products, and generate:
- a Markdown report in `output/`
- a DOCX version of the same report in `output/` (best-effort if python-docx is installed)

Constraints
-----------
- Rule-based parsing only (no external APIs).
- All output in English.
- Use emojis in terminal output and the report for scannability.
"""

from __future__ import annotations

import datetime as _dt
import os
import re
import sys
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, List, Optional, Tuple


# -----------------------------
# Data models (simple + explicit)
# -----------------------------


@dataclass
class Product:
    name: str
    list_price_eur: Optional[int] = None
    specs: Dict[str, str] = field(default_factory=dict)


@dataclass
class Lot:
    lot_id: str  # e.g. "Lot 1"
    title: str
    quantity: Optional[int] = None
    requirements: List[str] = field(default_factory=list)


@dataclass
class Tender:
    reference: Optional[str] = None
    buyer: Optional[str] = None
    title: Optional[str] = None
    publication_date: Optional[str] = None
    submission_deadline: Optional[str] = None
    primary_contact: Optional[str] = None
    scope_text: Optional[str] = None
    budget_text: Optional[str] = None
    lots: List[Lot] = field(default_factory=list)
    eligibility_criteria: List[str] = field(default_factory=list)
    required_documents: List[str] = field(default_factory=list)
    award_criteria_table: List[Tuple[str, str]] = field(default_factory=list)  # (criterion, weight)
    technical_value_breakdown: List[Tuple[str, str]] = field(default_factory=list)  # (item, weight)
    delivery_constraints: List[str] = field(default_factory=list)
    commercial_terms: List[str] = field(default_factory=list)
    penalties: List[str] = field(default_factory=list)
    variants: List[str] = field(default_factory=list)
    site_visit: List[str] = field(default_factory=list)


@dataclass
class CompanyProfile:
    name: str = "BrightEyes"
    founded_year: Optional[int] = None
    headquarters: Optional[str] = None
    employees: Optional[int] = None
    annual_revenue_eur: Optional[float] = None
    products: Dict[str, Product] = field(default_factory=dict)


@dataclass
class EligibilityResult:
    verdict: str  # GO / NO_GO / GO_WITH_RISKS
    met: List[str] = field(default_factory=list)
    risks: List[str] = field(default_factory=list)
    assumptions: List[str] = field(default_factory=list)


# -----------------------------
# Small utilities
# -----------------------------


def _read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def _ensure_dir(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)


def _money_to_number_eur(text: str) -> Optional[float]:
    """
    Best-effort parse of euro numbers like:
    - "€2.1M"
    - "€2,100,000"
    - "€500,000"
    Returns float (EUR) or None.
    """
    if not text:
        return None
    s = text.strip()
    s = s.replace("EUR", "€")
    m = re.search(r"€\s*([0-9]+(?:[.,][0-9]+)?)\s*([mMkK])?\b", s)
    if not m:
        # try plain number without €
        m2 = re.search(r"([0-9]{1,3}(?:[.,][0-9]{3})+|[0-9]+(?:[.,][0-9]+)?)", s)
        if not m2:
            return None
        num = m2.group(1)
        num = num.replace(",", "")
        try:
            return float(num)
        except ValueError:
            return None
    num_s = m.group(1).replace(",", "")
    try:
        val = float(num_s)
    except ValueError:
        return None
    suffix = (m.group(2) or "").lower()
    if suffix == "m":
        val *= 1_000_000
    elif suffix == "k":
        val *= 1_000
    return val


def _extract_first(pattern: str, text: str, flags: int = 0) -> Optional[str]:
    m = re.search(pattern, text, flags)
    return m.group(1).strip() if m else None


def _parse_simple_bullets(section_text: str) -> List[str]:
    """
    Extract bullet lines in Markdown starting with '-' from a section.
    Keeps the original text (already English in our demo tender).
    """
    out: List[str] = []
    for line in section_text.splitlines():
        line = line.strip()
        if line.startswith("- "):
            out.append(line[2:].strip())
    return out


# -----------------------------
# Parsing company context (rule-based)
# -----------------------------


def parse_company_context(md_text: str) -> CompanyProfile:
    profile = CompanyProfile()

    profile.headquarters = _extract_first(r"based in ([^.]+)\.", md_text, flags=re.IGNORECASE)

    founded = _extract_first(r"Founded in (\d{4})", md_text, flags=re.IGNORECASE)
    if founded and founded.isdigit():
        profile.founded_year = int(founded)

    employees = _extract_first(r"Currently\s+(\d+)\s+employees", md_text, flags=re.IGNORECASE)
    if employees and employees.isdigit():
        profile.employees = int(employees)

    # Revenue line like "- Revenue: €2.1M"
    rev_line = _extract_first(r"Revenue:\s*([^\n]+)", md_text, flags=re.IGNORECASE)
    if rev_line:
        profile.annual_revenue_eur = _money_to_number_eur(rev_line)

    # Products: headings "### <Product name>" followed by bullets and list price
    product_blocks = re.split(r"\n###\s+", md_text)
    for block in product_blocks[1:]:
        lines = block.splitlines()
        name = lines[0].strip()
        body = "\n".join(lines[1:])
        price_line = _extract_first(r"\*\*List price:\s*€\s*([0-9,]+)\*\*", body)
        price: Optional[int] = None
        if price_line:
            try:
                price = int(price_line.replace(",", "").strip())
            except ValueError:
                price = None
        specs: Dict[str, str] = {}
        for bullet in _parse_simple_bullets(body):
            # store as raw bullet text; also attempt a simple "key: value" split
            if ":" in bullet:
                k, v = bullet.split(":", 1)
                specs[k.strip()] = v.strip()
            else:
                specs[bullet] = "yes"
        profile.products[name] = Product(name=name, list_price_eur=price, specs=specs)

    return profile


# -----------------------------
# Parsing tender (rule-based)
# -----------------------------


def parse_tender(md_text: str) -> Tender:
    t = Tender()

    # Header fields (best-effort)
    t.reference = _extract_first(r"\*\*Reference:\*\*\s*([^\n]+)", md_text)
    t.publication_date = _extract_first(r"\*\*Publication date:\*\*\s*([^\n]+)", md_text)
    t.submission_deadline = _extract_first(r"\*\*Submission deadline:\*\*\s*([^\n]+)", md_text)
    t.primary_contact = _extract_first(r"\*\*Contact:\*\*\s*([^\n]+)", md_text)

    # Buyer + tender title from headings
    t.buyer = _extract_first(r"##\s+([^\n]+)\n###", md_text)
    t.title = _extract_first(r"###\s+([^\n]+)", md_text)

    # Scope section: grab paragraph after "## 1. SCOPE OF CONTRACT"
    scope = _extract_first(
        r"##\s*1\.\s*SCOPE OF CONTRACT\s*\n+([\s\S]*?)\n+\s*##\s*2\.\s*LOTS",
        md_text,
        flags=re.IGNORECASE,
    )
    if scope:
        t.scope_text = " ".join([ln.strip() for ln in scope.splitlines() if ln.strip()])

    # Lots: split at "### Lot X — Title"
    lot_iter = list(
        re.finditer(r"###\s+(Lot\s+\d+)\s+—\s+([^\n]+)\n([\s\S]*?)(?=\n###\s+Lot\s+\d+\s+—|\n##\s*3\.)", md_text)
    )
    for m in lot_iter:
        lot_id = m.group(1).strip()
        title = m.group(2).strip()
        body = m.group(3)
        qty = _extract_first(r"-\s*Quantity:\s*(\d+)\s*units?", body, flags=re.IGNORECASE)
        quantity = int(qty) if qty and qty.isdigit() else None
        requirements = _parse_simple_bullets(body)
        t.lots.append(Lot(lot_id=lot_id, title=title, quantity=quantity, requirements=requirements))

    # Eligibility criteria
    elig_block = _extract_first(
        r"###\s*3\.1\s*Eligibility Criteria\s*\n([\s\S]*?)\n\s*###\s*3\.2",
        md_text,
        flags=re.IGNORECASE,
    )
    if elig_block:
        t.eligibility_criteria = _parse_simple_bullets(elig_block)

    docs_block = _extract_first(
        r"###\s*3\.2\s*Required Documents\s*\n([\s\S]*?)\n\s*##\s*4\.",
        md_text,
        flags=re.IGNORECASE,
    )
    if docs_block:
        t.required_documents = _parse_simple_bullets(docs_block)

    # Award criteria table rows
    # Markdown table like "| Criterion | Weight |"
    table_block = _extract_first(
        r"##\s*4\.\s*AWARD CRITERIA\s*\n([\s\S]*?)\n\s*###\s*Technical value breakdown",
        md_text,
        flags=re.IGNORECASE,
    )
    if table_block:
        for line in table_block.splitlines():
            line = line.strip()
            if not line.startswith("|"):
                continue
            if "Criterion" in line and "Weight" in line:
                continue
            if set(line.replace("|", "").strip()) <= {"-", ":"}:
                continue
            cols = [c.strip() for c in line.strip("|").split("|")]
            if len(cols) >= 2 and cols[0] and cols[1]:
                t.award_criteria_table.append((cols[0], cols[1]))

    tech_breakdown_block = _extract_first(
        r"###\s*Technical value breakdown\s*\(50%\):\s*\n([\s\S]*?)\n\s*##\s*5\.",
        md_text,
        flags=re.IGNORECASE,
    )
    if tech_breakdown_block:
        # Lines like "- Performance and technical specifications: 20%"
        for bullet in _parse_simple_bullets(tech_breakdown_block):
            if ":" in bullet:
                k, v = bullet.split(":", 1)
                t.technical_value_breakdown.append((k.strip(), v.strip()))
            else:
                t.technical_value_breakdown.append((bullet.strip(), ""))

    # Budget + commercial terms
    financial_block = _extract_first(
        r"##\s*5\.\s*FINANCIAL CONDITIONS\s*\n([\s\S]*?)\n\s*##\s*6\.",
        md_text,
        flags=re.IGNORECASE,
    )
    if financial_block:
        bullets = _parse_simple_bullets(financial_block)
        t.commercial_terms.extend(bullets)
        for b in bullets:
            if "budget" in b.lower():
                t.budget_text = b

    delivery_block = _extract_first(
        r"##\s*6\.\s*DELIVERY TIMELINE\s*\n([\s\S]*?)\n\s*##\s*7\.",
        md_text,
        flags=re.IGNORECASE,
    )
    if delivery_block:
        t.delivery_constraints = _parse_simple_bullets(delivery_block)

    penalties_block = _extract_first(
        r"##\s*7\.\s*LATE DELIVERY PENALTIES\s*\n([\s\S]*?)\n\s*##\s*8\.",
        md_text,
        flags=re.IGNORECASE,
    )
    if penalties_block:
        t.penalties = _parse_simple_bullets(penalties_block)

    variants_block = _extract_first(
        r"##\s*8\.\s*VARIANTS\s*\n([\s\S]*?)\n\s*##\s*9\.",
        md_text,
        flags=re.IGNORECASE,
    )
    if variants_block:
        t.variants = _parse_simple_bullets(variants_block)

    site_visit_block = _extract_first(
        r"##\s*9\.\s*SITE VISIT\s*\n([\s\S]*?)\n\s*---",
        md_text,
        flags=re.IGNORECASE,
    )
    if site_visit_block:
        t.site_visit = [ln.strip() for ln in site_visit_block.splitlines() if ln.strip()]

    return t


# -----------------------------
# Eligibility check (rule-based)
# -----------------------------


def run_eligibility_check(tender: Tender, company: CompanyProfile) -> EligibilityResult:
    now_year = _dt.date.today().year
    res = EligibilityResult(verdict="GO_WITH_RISKS")

    # Experience ≥ 3 years
    exp_req_years = 3
    if company.founded_year:
        years = max(0, now_year - company.founded_year)
        if years >= exp_req_years:
            res.met.append(f"Experience: inferred {years} years since founding ({company.founded_year}) meets ≥ {exp_req_years} years.")
        else:
            res.risks.append(f"RISK: Tender requires ≥ {exp_req_years} years of experience. Inferred experience is {years} years.")
    else:
        res.risks.append("RISK: Tender requires ≥ 3 years of experience. Company founding year not found in context.")

    # Revenue ≥ €500,000
    revenue_req = 500_000
    if company.annual_revenue_eur is not None:
        if company.annual_revenue_eur >= revenue_req:
            res.met.append(f"Revenue: €{company.annual_revenue_eur:,.0f}/year meets ≥ €{revenue_req:,.0f}.")
        else:
            res.risks.append(
                f"RISK: Tender requires minimum annual revenue €{revenue_req:,.0f}. Company context indicates €{company.annual_revenue_eur:,.0f}."
            )
    else:
        res.risks.append("RISK: Tender requires minimum annual revenue €500,000. Revenue not found in company context.")

    # CE certification class IIa min (not in company context)
    res.risks.append("RISK: Tender requires CE medical marking (Class IIa minimum). Company context does not state CE class/certification evidence.")
    res.assumptions.append("Assumption: Proposed devices are CE-marked (Class IIa or higher); attach CE certificates per device.")

    # References: at least 3 similar installations in public hospitals in Europe within last 5 years
    res.risks.append("RISK: Tender requires ≥ 3 similar installations in public hospitals in Europe (last 5 years). Not evidenced in company context.")
    res.assumptions.append("Assumption: BrightEyes can provide ≥ 3 qualifying public-hospital references; collect reference letters + contact details.")

    # Maintenance within 48 hours
    res.risks.append("RISK: Tender requires on-site maintenance capability within 48 hours. Capability not evidenced in company context.")
    res.assumptions.append("Assumption: BrightEyes can commit to a 48h on-site SLA in Lombardy via internal technicians and/or certified partners.")

    # Verdict
    # If any "hard fail" signals exist (e.g. revenue below threshold), change verdict to NO_GO.
    hard_fail = any("Revenue:" in m and "meets" not in m for m in res.met)  # defensive (unused)
    if any("Inferred experience is" in r for r in res.risks) or any("indicates" in r and "requires" in r for r in res.risks):
        # If we found an explicit experience/revenue miss, treat as NO_GO
        if any("Inferred experience is" in r for r in res.risks) or any("indicates €" in r and "requires minimum annual revenue" in r for r in res.risks):
            res.verdict = "NO_GO"
        else:
            res.verdict = "GO_WITH_RISKS"
    else:
        # Even if all clear, keep GO_WITH_RISKS when CE/references/48h are unknown.
        res.verdict = "GO_WITH_RISKS" if res.risks else "GO"

    return res


# -----------------------------
# Product matching (rule-based)
# -----------------------------


def _get_product(company: CompanyProfile, contains: str) -> Optional[Product]:
    for name, p in company.products.items():
        if contains.lower() in name.lower():
            return p
    return None


def match_products(tender: Tender, company: CompanyProfile) -> Dict[str, Dict[str, object]]:
    """
    Returns: lot_id -> dict with recommended_product, fit_notes, gaps
    """
    matches: Dict[str, Dict[str, object]] = {}

    retina = _get_product(company, "RetinaScan")
    octp = _get_product(company, "OCT")
    autoref = _get_product(company, "AutoRef")

    for lot in tender.lots:
        lot_key = lot.lot_id
        rec: Optional[Product] = None
        fit: List[str] = []
        gaps: List[str] = []

        title_l = lot.title.lower()
        if "retinal" in title_l or "retina" in title_l:
            rec = retina
            if rec:
                # requirements checks (very lightweight string rules)
                if any("DICOM export" in r or "DICOM" in r for r in lot.requirements):
                    if any("DICOM" in k.upper() or "DICOM export" in k for k in rec.specs.keys()):
                        fit.append("DICOM export: supported (per company catalog).")
                    else:
                        gaps.append("RISK: DICOM export is mandatory; catalog does not explicitly list DICOM support for this product.")
                # Resolution
                if "14 megapixel" in " ".join(rec.specs.keys()).lower() or "14 megapixel" in " ".join(rec.specs.values()).lower():
                    fit.append("Resolution: 14MP meets minimum 12MP.")
                else:
                    fit.append("Resolution: appears compliant (verify ≥ 12MP).")
                # FOV
                if "60°" in " ".join(rec.specs.keys()) or "60°" in " ".join(rec.specs.values()):
                    fit.append("Field of capture: 60° matches 'ideally 60°' requirement.")
                else:
                    fit.append("Field of capture: appears compliant (verify ≥ 45° / ideally 60°).")
                # HIS integration
                if any("Dedalus ORBIS" in r for r in lot.requirements):
                    if any("HL7" in k.upper() or "FHIR" in k.upper() for k in rec.specs.keys()) or any(
                        "HL7" in v.upper() or "FHIR" in v.upper() for v in rec.specs.values()
                    ):
                        fit.append("HIS integration: HL7 FHIR compatibility suggests feasible Dedalus ORBIS integration (needs confirmation).")
                        gaps.append("RISK: Confirm Dedalus ORBIS integration specifics (interfaces, workflows, test plan).")
                    else:
                        gaps.append("RISK: Dedalus ORBIS compatibility required; no explicit interoperability spec found in catalog.")

        elif "oct" in title_l or "coherence" in title_l:
            rec = octp
            if rec:
                if any("DICOM" in r for r in lot.requirements):
                    if any("DICOM" in k.upper() for k in rec.specs.keys()) or any("DICOM" in v.upper() for v in rec.specs.values()):
                        fit.append("DICOM export: supported (per company catalog).")
                    else:
                        gaps.append("RISK: DICOM export is mandatory; catalog does not explicitly list DICOM support for this product.")
                # speed
                speed_blob = " ".join(rec.specs.keys()).lower() + " " + " ".join(rec.specs.values()).lower()
                if "100,000" in speed_blob or "100000" in speed_blob:
                    fit.append("Scan speed: 100,000 A-scans/sec meets ≥ 80,000 requirement.")
                else:
                    fit.append("Scan speed: appears compliant (verify ≥ 80,000 A-scans/sec).")
                if "anterior segment" in speed_blob:
                    fit.append("Anterior segment module: included (meets 'preferred').")
                else:
                    fit.append("Anterior segment module: not clearly stated (verify).")
                if "normative database" in speed_blob or "database" in speed_blob:
                    fit.append("Normative database: present (per catalog).")
                else:
                    gaps.append("RISK: Normative database required; not clearly stated in catalog extract.")
                if "ai" in speed_blob or "retinal layer" in speed_blob:
                    fit.append("Automatic retinal layer analysis / AI: present (per catalog; highlight for innovation points).")
                else:
                    fit.append("Automatic retinal layer analysis: verify feature availability and regulatory status.")

        elif "refract" in title_l or "autorefract" in title_l:
            rec = autoref
            if rec:
                blob_raw = " ".join(rec.specs.keys()) + " " + " ".join(rec.specs.values())
                blob = blob_raw.lower()
                if "keratometer" in blob:
                    fit.append("Integrated keratometer: supported (per catalog).")
                else:
                    gaps.append("RISK: Integrated keratometer required; not found in catalog extract.")
                if "3 seconds" in blob or "3 second" in blob or "3s" in blob:
                    fit.append("Measurement time: 3 seconds meets < 5 seconds.")
                else:
                    fit.append("Measurement time: appears compliant (verify < 5 seconds).")
                # UI language detection: company context lists "Multilingual interface (EN, FR, IT, DE, ES)"
                if re.search(r"\bitalian\b", blob_raw, flags=re.IGNORECASE) or re.search(r"\bIT\b", blob_raw):
                    fit.append("User interface: Italian supported (per catalog multilingual interface).")
                else:
                    gaps.append("RISK: UI in Italian required; not evidenced in catalog extract.")

        matches[lot_key] = {
            "recommended_product": rec,
            "fit_notes": fit,
            "gaps": gaps,
        }

    return matches


# -----------------------------
# Reporting
# -----------------------------


def _verdict_emoji(verdict: str) -> str:
    if verdict == "GO":
        return "🟢"
    if verdict == "GO_WITH_RISKS":
        return "🟡"
    return "🔴"


def generate_markdown_report(
    tender: Tender,
    company: CompanyProfile,
    eligibility: EligibilityResult,
    matches: Dict[str, Dict[str, object]],
    generated_at: _dt.datetime,
) -> str:
    # Summary fields
    budget = tender.budget_text or "Not found"
    lines: List[str] = []

    lines.append(f"# 📄 Tender Analysis — {tender.reference or 'UNKNOWN_REFERENCE'}")
    lines.append("")
    lines.append("## 🧾 Summary")
    lines.append("")
    lines.append(f"- **Buyer**: {tender.buyer or 'Not found'}")
    lines.append(f"- **Reference**: {tender.reference or 'Not found'}")
    lines.append(f"- **Publication date**: {tender.publication_date or 'Not found'}")
    lines.append(f"- **Submission deadline**: {tender.submission_deadline or 'Not found'}")
    lines.append(f"- **Budget (estimated)**: {budget}")
    lines.append(f"- **Primary contact**: {tender.primary_contact or 'Not found'}")
    lines.append("")

    lines.append("## 🎯 Scope (as stated)")
    lines.append("")
    lines.append(tender.scope_text or "Not found.")
    lines.append("")

    # Eligibility
    lines.append("## ✅ Eligibility — Go/No-Go")
    lines.append("")
    lines.append(f"**Verdict: {eligibility.verdict}** {_verdict_emoji(eligibility.verdict)}")
    lines.append("")

    if eligibility.met:
        lines.append("### ✅ Met")
        lines.append("")
        for item in eligibility.met:
            lines.append(f"- {item}")
        lines.append("")

    if eligibility.risks:
        lines.append("### ⚠️ RISK items (must be addressed in submission)")
        lines.append("")
        for item in eligibility.risks:
            lines.append(f"- {item}")
        lines.append("")

    if eligibility.assumptions:
        lines.append("### 📎 Assumptions / evidence to attach")
        lines.append("")
        for item in eligibility.assumptions:
            lines.append(f"- {item}")
        lines.append("")

    # Award criteria
    lines.append("## 🏅 Award Criteria (scoring weights)")
    lines.append("")
    if tender.award_criteria_table:
        lines.append("| Criterion | Weight |")
        lines.append("|---|---:|")
        for c, w in tender.award_criteria_table:
            lines.append(f"| {c} | {w} |")
    else:
        lines.append("Not found.")
    lines.append("")

    if tender.technical_value_breakdown:
        lines.append("### 🧠 Technical value breakdown")
        lines.append("")
        for k, v in tender.technical_value_breakdown:
            if v:
                lines.append(f"- {k} — {v}")
            else:
                lines.append(f"- {k}")
        lines.append("")

    # Lots + matching
    lines.append("## 🔗 Lots — Requirements & BrightEyes Product Fit")
    lines.append("")
    for lot in tender.lots:
        lines.append(f"### {lot.lot_id} — {lot.title}")
        lines.append("")
        if lot.quantity is not None:
            lines.append(f"- **Quantity**: {lot.quantity}")
        else:
            lines.append("- **Quantity**: Not found")
        lines.append("- **Key requirements**:")
        for req in lot.requirements:
            lines.append(f"  - {req}")
        lines.append("")

        m = matches.get(lot.lot_id, {})
        rec: Optional[Product] = m.get("recommended_product")  # type: ignore[assignment]
        if rec:
            price = f"€{rec.list_price_eur:,.0f}" if rec.list_price_eur else "price not found"
            lines.append(f"- **Recommended product**: {rec.name} (list price: {price})")
        else:
            lines.append("- **Recommended product**: Not found (no match rule triggered)")

        fit_notes: List[str] = m.get("fit_notes", [])  # type: ignore[assignment]
        gaps: List[str] = m.get("gaps", [])  # type: ignore[assignment]
        if fit_notes:
            lines.append("- **Fit notes**:")
            for n in fit_notes:
                lines.append(f"  - {n}")
        if gaps:
            lines.append("- **Gaps / RISK**:")
            for g in gaps:
                lines.append(f"  - {g}")
        lines.append("")

    # Deadlines & delivery
    lines.append("## 📅 Key Deadlines & Delivery Constraints")
    lines.append("")
    if tender.submission_deadline:
        lines.append(f"- **Submission deadline**: {tender.submission_deadline}")
    if tender.delivery_constraints:
        lines.append("- **Delivery / installation / training**:")
        for d in tender.delivery_constraints:
            lines.append(f"  - {d}")
    if tender.site_visit:
        lines.append("- **Site visit**:")
        for ln in tender.site_visit:
            # Keep as-is; already English in this tender
            lines.append(f"  - {ln}")
    lines.append("")

    # Commercial terms
    lines.append("## 💰 Commercial Terms (from tender)")
    lines.append("")
    if tender.commercial_terms:
        for term in tender.commercial_terms:
            lines.append(f"- {term}")
    else:
        lines.append("Not found.")
    lines.append("")

    if tender.penalties:
        lines.append("### ⏱️ Late delivery penalties")
        lines.append("")
        for p in tender.penalties:
            lines.append(f"- {p}")
        lines.append("")

    # Variants + strategy
    lines.append("## 🧩 Variants & Strategy Recommendations")
    lines.append("")
    if tender.variants:
        lines.append("- **Variants permitted (tender)**:")
        for v in tender.variants:
            lines.append(f"  - {v}")
        lines.append("")

    lines.append("- **Recommended focus (based on weights)**:")
    lines.append("  - Maximize **technical value (50%)** with a strong compliance matrix, interoperability plan (Dedalus ORBIS), and training program.")
    lines.append("  - Be price-competitive (30%) via **multi-lot packaging** and volume discounts (explicitly allowed).")
    lines.append("  - De-risk timeline (10%) with a detailed project plan that meets the 12-week delivery + 4-week commissioning constraints.")
    lines.append("  - Strengthen after-sales (10%) with a written **48h on-site SLA** and named local support resources.")
    lines.append("")

    # Draft outline
    lines.append("## 🧱 Draft Response Outline (skeleton)")
    lines.append("")
    lines.append("1. Administrative package (registration, insurance, CE certificates, references)")
    lines.append("2. Technical offer")
    lines.append("   - Lot-by-lot compliance matrix")
    lines.append("   - Interoperability & Dedalus ORBIS integration approach")
    lines.append("   - Training plan (agenda, duration, roles)")
    lines.append("   - Maintenance plan (SLA, preventive maintenance, escalation)")
    lines.append("3. Commercial offer")
    lines.append("   - Pricing per lot + optional multi-lot discount")
    lines.append("   - Payment terms acceptance")
    lines.append("   - Warranty + optional 3-year maintenance")
    lines.append("4. Delivery & installation plan")
    lines.append("5. Risk register and mitigations")
    lines.append("")

    lines.append("---")
    lines.append(f"Generated on {generated_at.strftime('%Y-%m-%d %H:%M')} by `tender_parser.py`")
    lines.append("")

    return "\n".join(lines)


def export_docx(markdown_text: str, docx_path: Path) -> Tuple[bool, str]:
    """
    Best-effort DOCX export.
    - If `python-docx` is installed, render Markdown as plain paragraphs (not full Markdown styling).
    - If not installed, returns (False, message).
    """
    try:
        from docx import Document  # type: ignore
    except Exception:
        return (
            False,
            "python-docx not installed. Install with `pip install python-docx` to enable DOCX export.",
        )

    doc = Document()
    for line in markdown_text.splitlines():
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        elif re.match(r"^\d+\.\s+", line):
            doc.add_paragraph(re.sub(r"^\d+\.\s+", "", line).strip(), style="List Number")
        else:
            doc.add_paragraph(line)
    doc.save(docx_path)
    return (True, "DOCX exported.")


# -----------------------------
# Main entrypoint
# -----------------------------


def main() -> int:
    root = Path(__file__).resolve().parent
    tender_path = root / "data" / "tender_hospital_milan.md"
    company_path = root.parent / "company_context.md"
    output_dir = root / "output"
    _ensure_dir(output_dir)

    print("📄 Loading tender and company context…")
    if not tender_path.exists():
        print(f"🔴 Tender file not found at: {tender_path}")
        return 2
    if not company_path.exists():
        print(f"🔴 Company context file not found at: {company_path}")
        return 2

    tender_md = _read_text(tender_path)
    company_md = _read_text(company_path)

    print("📋 Parsing tender…")
    tender = parse_tender(tender_md)
    print("🏢 Parsing company context…")
    company = parse_company_context(company_md)

    print("✅ Running eligibility check…")
    eligibility = run_eligibility_check(tender, company)

    print("🔗 Matching products to lots…")
    matches = match_products(tender, company)

    generated_at = _dt.datetime.now()
    report_md = generate_markdown_report(tender, company, eligibility, matches, generated_at=generated_at)

    ref_safe = re.sub(r"[^A-Za-z0-9_.-]+", "_", tender.reference or "UNKNOWN_REFERENCE")
    md_path = output_dir / f"tender_analysis_{ref_safe}.md"
    docx_path = output_dir / f"tender_analysis_{ref_safe}.docx"

    print(f"📝 Writing Markdown report to {md_path.relative_to(root)} …")
    md_path.write_text(report_md, encoding="utf-8")
    print("✅ Markdown saved.")

    print(f"📄 Exporting DOCX to {docx_path.relative_to(root)} …")
    ok, msg = export_docx(report_md, docx_path)
    if ok:
        print("✅ DOCX saved.")
    else:
        print(f"⚠️ DOCX skipped: {msg}")

    # Terminal summary (compact)
    print("\n📌 Summary")
    print(f"🏷️  Reference: {tender.reference or 'Not found'}")
    print(f"📅 Deadline: {tender.submission_deadline or 'Not found'}")
    print(f"{_verdict_emoji(eligibility.verdict)} Eligibility verdict: {eligibility.verdict}")
    if eligibility.risks:
        print(f"⚠️ Risks flagged: {len(eligibility.risks)}")
    print("✅ Done.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

